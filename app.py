import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image
import io
import datetime
import os
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER

st.set_page_config(page_title="Report Cover Page Generator", layout="centered")
st.title("📄Report Cover Page Generator")

# Sidebar Configuration
st.sidebar.header("Configuration")
prof_count = st.sidebar.selectbox("Number of Professors", [1, 2, 3])
student_count = st.sidebar.slider("Number of Students", min_value=1, max_value=10, value=1)

# University Selection (OUTSIDE FORM)
st.subheader("🎓 Select University")
predefined_universities = [
    "Ahsanullah University of Science and Technology",
    "American International University-Bangladesh",
    "Bangabandhu Sheikh Mujib Medical University",
    "Bangabandhu Sheikh Mujibur Rahman Agricultural University",
    "Bangladesh Agricultural University",
    "Bangladesh Open University",
    "Bangladesh University",
    "Bangladesh University of Business and Technology",
    "Bangladesh University of Engineering and Technology",
    "Bangladesh University of Health Sciences",
    "Bangladesh University of Professionals",
    "Begum Rokeya University",
    "Brac University",
    "Chittagong University",
    "Chittagong University of Engineering & Technology",
    "Chittagong Veterinary and Animal Sciences University",
    "Comilla University",
    "Daffodil International University",
    "East West University",
    "Hajee Mohammad Danesh Science and Technology University",
    "International Islamic University Chittagong",
    "Islamic University, Bangladesh",
    "Jagannath University",
    "Jahangirnagar University",
    "Jatiya Kabi Kazi Nazrul Islam University",
    "Khulna University",
    "Khulna University of Engineering & Technology",
    "Leading University",
    "North South University",
    "Pabna University of Science and Technology",
    "Patuakhali Science and Technology University",
    "Rajshahi University",
    "Rajshahi University of Engineering & Technology",
    "Shahjalal University of Science and Technology",
    "Sher-e-Bangla Agricultural University",
    "Sylhet Agricultural University",
    "Sylhet International University",
    "University of Asia Pacific",
    "University of Barishal",
    "University of Development Alternative",
    "University of Dhaka",
    "University of Information Technology & Sciences",
    "University of Rajshahi",
    "University of Science and Technology Chittagong",
    "Victoria University of Bangladesh",
    "Others"
]
selected_uni = st.selectbox("Choose University", predefined_universities)

if selected_uni == "Others":
    university_name = st.text_input("Enter University Name")
    logo = st.file_uploader("Upload University Logo", type=["png", "jpg", "jpeg"])
else:
    university_name = selected_uni
    image_name = selected_uni  # font related prob, for image name 
    logo_path = f"images/{image_name}.png"

    if os.path.exists(logo_path):
        logo = open(logo_path, "rb")
    else:
        st.warning(f"⚠️ Logo not found. Please upload it manually.")
        logo = st.file_uploader("Upload University Logo", type=["png", "jpg", "jpeg"])


# Cover Form
with st.form("cover_form"):
    st.subheader("📝 Fill in the details")

    department_name = st.text_input("Department Name")
    subject_code = st.text_input("Subject Code (e.g., CSE-2201)")
    subject_name = st.text_input("Subject Name")
    exp_no = st.text_input("Experiment No")
    exp_name = st.text_input("Experiment Name")

    professors = []
    st.markdown("### 👨‍🏫 Professor Info")
    for i in range(prof_count):
        st.markdown(f"**Professor {i+1}**")
        name = st.text_input("Name", key=f"prof_name_{i}")
        desig = st.text_input("Designation", key=f"prof_desig_{i}")
        professors.append((name, desig))

    students = []
    st.markdown("### 👨‍🎓 Student Info")
    for i in range(student_count):
        st.markdown(f"**Student {i+1}**")
        name = st.text_input("Student Name", key=f"stu_name_{i}")
        sid = st.text_input("Student ID", key=f"stu_id_{i}")
        students.append((name, sid))

    experiment_date = st.date_input("Experiment Date", value=datetime.date.today())
    submitted = st.form_submit_button("Generate Cover Page")

# Helper function to add centered paragraph in DOCX
def add_centered_paragraph(doc, text, font_size=14, bold=False, space_before=0, space_after=0.8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p

# Generate DOCX
def generate_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    if logo:
        image = Image.open(logo)
        image_stream = io.BytesIO()
        image.save(image_stream, format="PNG")
        image_stream.seek(0)
        doc.add_picture(image_stream, width=Inches(1.2))
        doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    add_centered_paragraph(doc, university_name, font_size=18, bold=True, space_after=2)
    add_centered_paragraph(doc, department_name, font_size=16, space_after=6)
    add_centered_paragraph(doc, f"{subject_code} : {subject_name}", font_size=14, space_after=4)
    add_centered_paragraph(doc, f"Experiment No: {exp_no}", font_size=14)
    add_centered_paragraph(doc, f"Experiment Name: {exp_name}", font_size=14, bold=True)

    if any(name.strip() for name, _ in professors):
        add_centered_paragraph(doc, "", space_after=4)
        add_centered_paragraph(doc, "Submitted to:", font_size=16, bold=True, space_after=3)
        for name, desig in professors:
            if name:
                add_centered_paragraph(doc, name, font_size=14)
            if desig:
                add_centered_paragraph(doc, desig, font_size=11)

    if any(name.strip() for name, _ in students):
        add_centered_paragraph(doc, "", space_after=4)
        add_centered_paragraph(doc, "Submitted by:", font_size=16, bold=True, space_after=3)
        for name, sid in students:
            if name:
                add_centered_paragraph(doc, name, font_size=14)
            if sid:
                add_centered_paragraph(doc, f"ID: {sid}", font_size=14)

    add_centered_paragraph(doc, "", space_after=3)
    add_centered_paragraph(doc, f"Date of Experiment: {experiment_date.strftime('%B %d, %Y')}", font_size=14)

    return doc

# Generate PDF using ReportLab
def generate_pdf():
    pdf_buffer = io.BytesIO()
    doc_pdf = SimpleDocTemplate(pdf_buffer, pagesize=letter,
                                rightMargin=22, leftMargin=22,
                                topMargin=22, bottomMargin=22)

    styles = getSampleStyleSheet()

    style_title = ParagraphStyle(
        'Title',
        parent=styles['Title'],
        fontSize=18,
        alignment=TA_CENTER,
        spaceAfter=10,
        leading=20
    )
    style_heading = ParagraphStyle(
        'Heading2',
        parent=styles['Heading2'],
        fontSize=16,
        alignment=TA_CENTER,
        spaceAfter=10,
        leading=18
    )
    style_normal = ParagraphStyle(
        'Normal',
        parent=styles['Normal'],
        fontSize=14,
        alignment=TA_CENTER,
        spaceAfter=8,
        leading=16
    )

    story = []

    if logo:
        image = Image.open(logo)
        img_byte_arr = io.BytesIO()
        image.save(img_byte_arr, format='PNG')
        img_byte_arr.seek(0)
        rl_image = RLImage(img_byte_arr, width=1 * inch, height=1 * inch)
        rl_image.hAlign = 'CENTER'
        story.append(rl_image)
        story.append(Spacer(1, 15))

    story.append(Paragraph(university_name, style_title))
    story.append(Paragraph(department_name, style_heading))
    story.append(Spacer(1, 15))
    story.append(Paragraph(f"{subject_code} : {subject_name}", style_normal))
    story.append(Paragraph(f"Experiment No: {exp_no}", style_normal))
    story.append(Paragraph(f"Experiment Name: <b>{exp_name}</b>", style_normal))

    if any(name.strip() for name, _ in professors):
        story.append(Spacer(1, 15))
        story.append(Paragraph("Submitted to:", style_heading))
        for name, desig in professors:
            if name:
                story.append(Paragraph(name, style_normal))
            if desig:
                story.append(Paragraph(desig, style_normal))

    if any(name.strip() for name, _ in students):
        story.append(Spacer(1, 15))
        story.append(Paragraph("Submitted by:", style_heading))
        for name, sid in students:
            if name:
                story.append(Paragraph(name, style_normal))
            if sid:
                story.append(Paragraph(f"ID: {sid}", style_normal))

    story.append(Spacer(1, 15))
    story.append(Paragraph(f"Date of Experiment: {experiment_date.strftime('%B %d, %Y')}", style_normal))

    doc_pdf.build(story)
    pdf_buffer.seek(0)
    return pdf_buffer

# Select download format
format_choice = st.selectbox("Select download format", ["DOCX", "PDF"])

# Handle Submission
if submitted:
    if format_choice == "DOCX":
        doc = generate_docx()
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.success("✅ DOCX Cover page generated successfully!")
        st.download_button(
            label="📄 Download DOCX",
            data=buffer,
            file_name="Lab_Report_Cover_Page.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        try:
            pdf_buffer = generate_pdf()
            st.success("✅ PDF Cover page generated successfully!")
            st.download_button(
                label="📄 Download PDF",
                data=pdf_buffer,
                file_name="Lab_Report_Cover_Page.pdf",
                mime="application/pdf"
            )
        except Exception as e:
            st.error("❌ PDF generation failed.")
            st.error(str(e))
